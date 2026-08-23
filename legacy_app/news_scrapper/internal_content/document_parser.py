"""Server-side document validation and text extraction for contributions.

PDFs are parsed with pypdf and DOCX files with python-docx. Uploads are
validated by extension, declared type, and binary signature so a renamed
arbitrary file cannot pass. Every failure raises ContributionError with a
message that is safe to show directly in the UI; technical detail goes to the
application log instead.
"""

from __future__ import annotations

import io
import logging
import re
import zipfile

import docx
from pypdf import PdfReader

logger = logging.getLogger(__name__)

DOCUMENT_MAX_BYTES = 25 * 1024 * 1024
PDF_MAX_PAGES = 100
EXTRACT_MAX_CHARS = 200_000
MIN_READABLE_CHARS = 40

WORD_DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_TYPE = "application/pdf"

_BLOCKED_EXTENSION_MESSAGES = {
    "doc": "This Word format is not supported yet. Save the document as .docx and try again.",
}
_BLOCKED_GENERIC_EXTENSIONS = {
    "docm", "xls", "xlsx", "ppt", "pptx",
    "zip", "rar", "7z", "tar", "gz",
    "exe", "dll", "bat", "cmd", "sh", "msi",
}


class ContributionError(Exception):
    """A user-facing failure. The message may be shown directly in the UI."""


def extension_of(filename: str) -> str:
    name = str(filename or "")
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


def validate_document_upload(filename: str, declared_type: str, size_bytes: int) -> str:
    """Return 'pdf' or 'docx' for an acceptable upload; raise otherwise."""

    extension = extension_of(filename)
    if extension == "doc":
        raise ContributionError(_BLOCKED_EXTENSION_MESSAGES["doc"])
    if extension in _BLOCKED_GENERIC_EXTENSIONS:
        raise ContributionError(
            f"{filename} is not a supported format. Choose a PDF or Word (.docx) document."
        )
    is_pdf = extension == "pdf" or declared_type == PDF_TYPE
    is_docx = extension == "docx" or declared_type == WORD_DOCX_TYPE
    if not (is_pdf or is_docx):
        raise ContributionError("Choose a PDF or Word (.docx) document to import.")
    if size_bytes > DOCUMENT_MAX_BYTES:
        readable = DOCUMENT_MAX_BYTES / (1024 * 1024)
        raise ContributionError(
            f"{filename} is larger than the {readable:.0f} MB document limit. "
            "Split it or export a smaller copy."
        )
    return "pdf" if is_pdf else "docx"


def _signature_matches(kind: str, data: bytes) -> bool:
    if kind == "pdf":
        return data[:5] == b"%PDF-"
    if kind == "docx":
        if data[:4] != b"PK\x03\x04":
            return False
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                names = archive.namelist()
            return any(name.startswith("word/") for name in names)
        except zipfile.BadZipFile:
            return False
    return False


def _normalize_extracted(value: str) -> str:
    cleaned = str(value or "")
    cleaned = cleaned.replace("\u00a0", " ")
    cleaned = re.sub(r"([A-Za-z])-[ \t]*\r?\n[ \t]*([a-z])", r"\1\2", cleaned)
    cleaned = re.sub(r"[ \t]+\r?\n", "\n", cleaned)
    cleaned = re.sub(r"\r?\n[ \t]+", "\n", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def detect_title(text: str, filename: str) -> str:
    """Deterministic title suggestion: first meaningful line, else filename."""

    for line in (part.strip() for part in str(text or "").splitlines()):
        if 3 <= len(line) <= 120:
            return line
    base = re.sub(r"\.(pdf|docx)$", "", str(filename or ""), flags=re.IGNORECASE)
    base = re.sub(r"[_-]+", " ", base).strip()
    return base[:120]


def _empty_text_error(kind: str) -> ContributionError:
    if kind == "pdf":
        return ContributionError(
            "We couldn't find readable text in this PDF. It may contain scanned pages "
            "rather than selectable text. Upload a text-based PDF/DOCX or start a story manually."
        )
    return ContributionError(
        "We couldn't find readable text in this document. Upload a text-based PDF or "
        "DOCX, or start a story manually."
    )


def extract_pdf(data: bytes, filename: str) -> tuple[str, int]:
    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as error:  # noqa: BLE001 - parser raises many internal types
        logger.warning("[internal-content] PDF open failed for %s: %s", filename, error)
        raise ContributionError(
            f"{filename} could not be opened as a PDF. The file may be damaged."
        ) from error

    if reader.is_encrypted:
        raise ContributionError(
            "This PDF is password protected. Remove the password or upload an unprotected copy."
        )

    page_count = len(reader.pages)
    if page_count > PDF_MAX_PAGES:
        raise ContributionError(
            f"This PDF has {page_count} pages and the import limit is {PDF_MAX_PAGES}. "
            "Split the document into smaller parts and import them separately."
        )

    pages: list[str] = []
    try:
        for page in reader.pages:
            pages.append(_normalize_extracted(page.extract_text() or ""))
    except Exception as error:  # noqa: BLE001
        logger.warning("[internal-content] PDF extraction failed for %s: %s", filename, error)
        raise ContributionError(
            f"{filename} could not be read as a PDF. It may be damaged or use an unsupported layout."
        ) from error

    text = _normalize_extracted("\n\n".join(pages))
    return text, page_count


def extract_docx(data: bytes, filename: str) -> tuple[str, int]:
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as error:  # noqa: BLE001
        logger.warning("[internal-content] DOCX open failed for %s: %s", filename, error)
        raise ContributionError(
            f"{filename} could not be read as a Word document. It may be damaged."
        ) from error

    blocks: list[str] = [
        paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
    ]
    for table in document.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        blocks.extend(row for row in rows if row.strip())

    text = _normalize_extracted("\n\n".join(blocks))
    return text, None


def parse_document(filename: str, declared_type: str, data: bytes) -> dict:
    """Validate and fully extract one upload.

    Returns {text, page_count, character_count, detected_title, paragraphs}.
    """

    size_bytes = len(data)
    kind = validate_document_upload(filename, declared_type, size_bytes)

    if not _signature_matches(kind, data):
        expected = "a real PDF" if kind == "pdf" else "a Word .docx file"
        raise ContributionError(f"{filename} does not look like {expected}. Check the file and try again.")

    if kind == "pdf":
        text, page_count = extract_pdf(data, filename)
    else:
        text, page_count = extract_docx(data, filename)

    if len(text.replace("\n", "").replace(" ", "")) < MIN_READABLE_CHARS:
        raise _empty_text_error(kind)
    if len(text) > EXTRACT_MAX_CHARS:
        thousands = len(text) // 1000
        raise ContributionError(
            f"The extracted text is about {thousands}k characters, above the "
            f"{EXTRACT_MAX_CHARS:,} character import limit. Split the document and "
            "import the sections you need."
        )

    return {
        "text": text,
        "page_count": page_count,
        "character_count": len(text),
        "detected_title": detect_title(text, filename),
        "paragraphs": [part for part in text.split("\n\n") if part],
    }
