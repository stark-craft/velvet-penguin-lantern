"""End-to-end tests for the Internal Contribution feature.

The suite drives the real composition root through the ASGI harness with the
feature's runtime redirected into a temporary directory, so contributions,
originals, and covers never touch real user state.
"""

import io
import os
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

import docx
from PIL import Image
from pypdf import PdfReader, PdfWriter

import main as composition
from news_scrapper.internal_content import document_parser, image_processor
from news_scrapper.internal_content import access as internal_access
from news_scrapper.internal_content import storage as internal_storage
from tests.asgi_harness import request as asgi_request


def build_pdf(page_texts):
    """Assemble a minimal but fully valid multi-page PDF with text."""

    catalog_num = 1
    pages_num = 2
    page_nums = []
    content_nums = []
    next_num = 3
    for _ in page_texts:
        page_nums.append(next_num)
        content_nums.append(next_num + 1)
        next_num += 2
    font_num = next_num

    objects = {
        catalog_num: f"<< /Type /Catalog /Pages {pages_num} 0 R >>",
        pages_num: "<< /Type /Pages /Kids ["
        + " ".join(f"{num} 0 R" for num in page_nums)
        + f"] /Count {len(page_nums)} >>",
        font_num: "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    for text, page_num, content_num in zip(page_texts, page_nums, content_nums):
        objects[page_num] = (
            f"<< /Type /Page /Parent {pages_num} 0 R /MediaBox [0 0 612 792] "
            f"/Contents {content_num} 0 R /Resources << /Font << /F1 {font_num} 0 R >> >> >>"
        )
        objects[content_num] = {"stream": f"BT /F1 24 Tf 72 700 Td ({text}) Tj ET"}

    buffer = io.BytesIO()
    offsets = {}

    def write(chunk):
        buffer.write(chunk.encode("latin-1") if isinstance(chunk, str) else chunk)

    write("%PDF-1.4\n")
    for num in sorted(objects):
        offsets[num] = buffer.tell()
        value = objects[num]
        if isinstance(value, dict) and "stream" in value:
            payload = value["stream"].encode("latin-1")
            write(f"{num} 0 obj\n<< /Length {len(payload)} >>\nstream\n")
            write(payload)
            write("\nendstream\nendobj\n")
        else:
            write(f"{num} 0 obj\n{value}\nendobj\n")
    xref_offset = buffer.tell()
    count = max(objects) + 1
    write(f"xref\n0 {count}\n0000000000 65535 f \n")
    for num in range(1, count):
        write(f"{offsets.get(num, 0):010d} 00000 n \n")
    write(
        f"trailer\n<< /Size {count} /Root {catalog_num} 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    )
    return buffer.getvalue()


def build_encrypted_pdf():
    reader = PdfReader(io.BytesIO(build_pdf(["Confidential strategy"])))
    writer = PdfWriter()
    writer.append(reader)
    writer.encrypt("secret-passphrase")
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def build_docx(lines):
    document = docx.Document()
    for line in lines:
        document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_image(width, height, top_color=(210, 40, 40), bottom_color=None):
    image = Image.new("RGB", (width, height), top_color)
    if bottom_color is not None:
        for y in range(height // 2, height):
            for x in range(width):
                image.putpixel((x, y), bottom_color)
    buffer = io.BytesIO()
    image.save(buffer, "PNG")
    return buffer.getvalue()


def multipart_body(fields):
    boundary = "icformboundary736f6c7665"
    chunks = []
    for name, (filename, content, content_type) in fields.items():
        if isinstance(content, str):
            content = content.encode("utf-8")
        disposition = f'Content-Disposition: form-data; name="{name}"'
        if filename is not None:
            disposition += f'; filename="{filename}"'
        chunks.append(f"--{boundary}\r\n".encode())
        header = disposition + "\r\n"
        if content_type:
            header += f"Content-Type: {content_type}\r\n"
        chunks.append(header.encode("latin-1"))
        chunks.append(b"\r\n")
        chunks.append(content)
        chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode())
    return b"".join(chunks), f"multipart/form-data; boundary={boundary}"


class ApiClient:
    """ASGI harness client that keeps its signed viewer cookie across calls."""

    def __init__(self):
        self.cookies = {}

    def request(self, method, path, *, headers=None, json_body=None, body=None):
        sent_headers = dict(headers or {})
        if self.cookies:
            sent_headers.setdefault(
                "Cookie",
                "; ".join(f"{name}={value}" for name, value in self.cookies.items()),
            )
        response = asgi_request(
            composition.app,
            method,
            path,
            headers=sent_headers,
            json_body=json_body,
            body=body,
        )
        set_cookie = response.headers.get("set-cookie")
        if set_cookie and "=" in set_cookie:
            pair = set_cookie.split(";", 1)[0]
            name, _, value = pair.partition("=")
            self.cookies[name.strip()] = value.strip()
        return response

    def upload(self, method, path, fields):
        payload, content_type = multipart_body(fields)
        return self.request(method, path, body=payload, headers={"Content-Type": content_type})


class InternalContentTests(unittest.TestCase):
    def setUp(self):
        self._temporary = tempfile.TemporaryDirectory()
        self.addCleanup(self._temporary.cleanup)
        runtime = Path(self._temporary.name) / "internal_content"
        for attribute, target in (
            ("RUNTIME_DIR", runtime),
            ("ORIGINALS_DIR", runtime / "originals"),
            ("COVERS_DIR", runtime / "covers"),
            ("CONTRIBUTIONS_FILE", runtime / "contributions.json"),
            ("NOTIFICATIONS_FILE", runtime / "internal_notifications.json"),
        ):
            patcher = patch.object(internal_storage, attribute, target)
            patcher.start()
            self.addCleanup(patcher.stop)
        access_patcher = patch.object(internal_access, "CONTRIBUTIONS_ALLOWED_IPS", {"10.0.0.25"})
        access_patcher.start()
        self.addCleanup(access_patcher.stop)

    # -- drafts -----------------------------------------------------------

    def test_manual_draft_create_update_and_delete(self):
        owner = ApiClient()
        created = owner.request("POST", "/internal-content/drafts", json_body={
            "title": "Display roadmap",
            "body": "First paragraph.\n\nSecond paragraph.",
            "category": "Technology",
            "team": "Display Research",
            "author": "Vineet",
        })
        self.assertEqual(created.status_code, 200)
        record = created.json()
        self.assertEqual(record["status"], "draft")
        self.assertEqual(record["content_type"], "story")
        self.assertEqual(record["title"], "Display roadmap")
        self.assertTrue(record["id"])

        updated = owner.request("PUT", f"/internal-content/{record['id']}", json_body={
            "title": "Display roadmap 2026",
            "summary": "A short summary.",
            "body": "Rewritten body.",
        })
        self.assertEqual(updated.status_code, 200)
        self.assertEqual(updated.json()["title"], "Display roadmap 2026")

        listed = owner.request("GET", "/internal-content/mine")
        self.assertEqual([item["id"] for item in listed.json()["items"]], [record["id"]])

        deleted = owner.request("DELETE", f"/internal-content/{record['id']}")
        self.assertEqual(deleted.status_code, 200)
        self.assertEqual(owner.request("GET", "/internal-content/mine").json()["items"], [])

    def test_owner_isolation_hides_other_viewers_records(self):
        owner = ApiClient()
        outsider = ApiClient()
        record = owner.request("POST", "/internal-content/drafts", json_body={
            "title": "Private draft",
            "body": "Owner-only words.",
        }).json()

        self.assertEqual(outsider.request("GET", f"/internal-content/{record['id']}").status_code, 404)
        self.assertEqual(outsider.request("GET", "/internal-content/mine").json()["items"], [])
        self.assertEqual(
            outsider.request("PUT", f"/internal-content/{record['id']}", json_body={"title": "stolen"}).status_code,
            404,
        )
        self.assertEqual(outsider.request("DELETE", f"/internal-content/{record['id']}").status_code, 404)

    # -- import -----------------------------------------------------------

    def test_import_creates_editable_draft_with_provenance(self):
        owner = ApiClient()
        pdf_bytes = build_pdf(["Quarterly display strategy", "Second page of evidence."])
        response = owner.upload("POST", "/internal-content/import", {
            "document": ("report.pdf", pdf_bytes, "application/pdf"),
        })
        self.assertEqual(response.status_code, 200)
        record = response.json()
        self.assertEqual(record["content_type"], "document_import")
        self.assertEqual(record["status"], "draft")
        self.assertIn("Quarterly display strategy", record["body"])
        source = record["source_document"]
        self.assertEqual(source["page_count"], 2)
        self.assertGreater(source["extracted_characters"], 10)
        stored_originals = list(Path(internal_storage.ORIGINALS_DIR).glob("*"))
        self.assertEqual(len(stored_originals), 1)
        self.assertNotEqual(stored_originals[0].name, "report.pdf")

        served = owner.request("GET", f"/internal-content/{record['id']}/document")
        self.assertEqual(served.status_code, 200)
        self.assertEqual(served.content, pdf_bytes)

    def test_document_validation_rejects_blocked_and_fake_files(self):
        owner = ApiClient()

        legacy_doc = owner.upload("POST", "/internal-content/import", {
            "document": ("memo.doc", b"old word binary", "application/msword"),
        })
        self.assertEqual(legacy_doc.status_code, 400)
        self.assertEqual(
            legacy_doc.json()["detail"],
            "This Word format is not supported yet. Save the document as .docx and try again.",
        )

        spreadsheet = owner.upload("POST", "/internal-content/import", {
            "document": ("sheet.xlsx", b"PK\x03\x04fake", ""),
        })
        self.assertEqual(spreadsheet.status_code, 400)

        fake_pdf = owner.upload("POST", "/internal-content/import", {
            "document": ("fake.pdf", b"this is not a pdf at all", "application/pdf"),
        })
        self.assertEqual(fake_pdf.status_code, 400)
        self.assertIn("does not look like a real PDF", fake_pdf.json()["detail"])

        fake_docx = owner.upload("POST", "/internal-content/import", {
            "document": ("fake.docx", b"PK\x03\x04 not really a docx package", ""),
        })
        self.assertEqual(fake_docx.status_code, 400)

    def test_oversize_document_is_rejected_before_parsing(self):
        owner = ApiClient()
        with patch.object(document_parser, "DOCUMENT_MAX_BYTES", 64):
            response = owner.upload("POST", "/internal-content/import", {
                "document": ("big.pdf", build_pdf(["A generous amount of real text."]), "application/pdf"),
            })
        self.assertEqual(response.status_code, 400)
        self.assertIn("document limit", response.json()["detail"])

    def test_encrypted_pdf_is_rejected_with_friendly_message(self):
        owner = ApiClient()
        response = owner.upload("POST", "/internal-content/import", {
            "document": ("locked.pdf", build_encrypted_pdf(), "application/pdf"),
        })
        self.assertEqual(response.status_code, 400)
        self.assertIn("password protected", response.json()["detail"])

    def test_scanned_or_empty_pdf_returns_readable_text_guidance(self):
        owner = ApiClient()
        # A page whose only text operator carries no glyphs extracts to no text.
        blank_page = build_pdf([""])
        response = owner.upload("POST", "/internal-content/import", {
            "document": ("scan.pdf", blank_page, "application/pdf"),
        })
        self.assertEqual(response.status_code, 400)
        detail = response.json()["detail"]
        self.assertIn("readable text", detail)
        self.assertIn("scanned pages", detail)

    def test_pdf_page_cap_rejects_over_100_pages(self):
        owner = ApiClient()
        many_pages = build_pdf([f"Page {index}" for index in range(101)])
        response = owner.upload("POST", "/internal-content/import", {
            "document": ("book.pdf", many_pages, "application/pdf"),
        })
        self.assertEqual(response.status_code, 400)
        self.assertIn("101 pages", response.json()["detail"])

    def test_extraction_cap_rejects_over_200k_characters(self):
        owner = ApiClient()
        huge_text = "x" * 205_000
        response = owner.upload("POST", "/internal-content/import", {
            "document": ("huge.pdf", build_pdf([huge_text]), "application/pdf"),
        })
        self.assertEqual(response.status_code, 400)
        self.assertIn("200,000 character import limit", response.json()["detail"])

    def test_docx_import_extracts_paragraph_and_table_text(self):
        owner = ApiClient()
        document = docx.Document()
        document.add_paragraph("Internal mobility program")
        document.add_paragraph("The program opens next quarter.")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Team"
        table.rows[0].cells[1].text = "Seats"
        buffer = io.BytesIO()
        document.save(buffer)
        response = owner.upload("POST", "/internal-content/import", {
            "document": ("notes.docx", buffer.getvalue(), ""),
        })
        self.assertEqual(response.status_code, 200)
        body = response.json()["body"]
        self.assertIn("Internal mobility program", body)
        self.assertIn("Team | Seats", body)
        self.assertIsNone(response.json()["source_document"]["page_count"])

    # -- covers -----------------------------------------------------------

    def upload_cover(self, client, record_id, image_bytes, focal_x=0.5, focal_y=0.5, name="cover.png"):
        return client.upload("POST", f"/internal-content/{record_id}/cover", {
            "cover": (name, image_bytes, "image/png"),
            "focal_x": (None, str(focal_x), None),
            "focal_y": (None, str(focal_y), None),
        })

    def create_draft(self, client, **fields):
        defaults = {"title": "Covered story", "body": "A body long enough to submit safely."}
        defaults.update(fields)
        return client.request("POST", "/internal-content/drafts", json_body=defaults).json()

    def test_cover_normalizes_to_exact_1600x900_webp(self):
        owner = ApiClient()
        record = self.create_draft(owner)
        response = self.upload_cover(owner, record["id"], build_image(1600, 900))
        self.assertEqual(response.status_code, 200)
        cover = response.json()["cover"]
        self.assertEqual(cover["width"], 1600)
        self.assertEqual(cover["height"], 900)

        served = owner.request("GET", f"/internal-content/{record['id']}/cover")
        self.assertEqual(served.status_code, 200)
        self.assertIn("image/webp", served.headers.get("content-type", ""))
        normalized = Image.open(io.BytesIO(served.content))
        self.assertEqual(normalized.format, "WEBP")
        self.assertEqual(normalized.size, (1600, 900))

    def test_portrait_cover_allowed_and_focal_point_changes_crop(self):
        owner = ApiClient()
        record = self.create_draft(owner)
        portrait = build_image(1080, 2400, top_color=(210, 30, 30), bottom_color=(20, 40, 220))

        top_upload = self.upload_cover(owner, record["id"], portrait, focal_x=0.5, focal_y=0.0)
        self.assertEqual(top_upload.status_code, 200)
        top_served = owner.request("GET", f"/internal-content/{record['id']}/cover")

        bottom_upload = self.upload_cover(owner, record["id"], portrait, focal_x=0.5, focal_y=1.0)
        self.assertEqual(bottom_upload.status_code, 200)
        bottom_served = owner.request("GET", f"/internal-content/{record['id']}/cover")

        self.assertNotEqual(top_served.content, bottom_served.content)

        top_image = Image.open(io.BytesIO(top_served.content)).convert("RGB")
        bottom_image = Image.open(io.BytesIO(bottom_served.content)).convert("RGB")
        self.assertEqual(top_image.size, (1600, 900))
        red = top_image.getpixel((800, 450))
        blue = bottom_image.getpixel((800, 450))
        self.assertGreater(red[0], 150)
        self.assertGreater(blue[2], 150)
        self.assertLess(red[2], blue[2])

        # Replacing the cover retires the previous normalized file.
        covers_after_replace = list(Path(internal_storage.COVERS_DIR).glob("*"))
        self.assertEqual(len(covers_after_replace), 1)

    def test_cover_validation_failures(self):
        owner = ApiClient()
        record = self.create_draft(owner)

        tiny = self.upload_cover(owner, record["id"], build_image(500, 500))
        self.assertEqual(tiny.status_code, 400)
        self.assertIn("960 x 540", tiny.json()["detail"])

        invalid = self.upload_cover(owner, record["id"], b"definitely not an image")
        self.assertEqual(invalid.status_code, 400)
        self.assertIn("could not be read", invalid.json()["detail"])

        gif_buffer = io.BytesIO()
        Image.new("RGB", (1200, 700), (10, 120, 10)).save(gif_buffer, "GIF")
        gif = self.upload_cover(owner, record["id"], gif_buffer.getvalue(), name="animated.gif")
        self.assertEqual(gif.status_code, 400)
        self.assertIn("GIF", gif.json()["detail"])

        with patch.object(image_processor, "COVER_MAX_BYTES", 128):
            oversized = self.upload_cover(owner, record["id"], build_image(1600, 900))
        self.assertEqual(oversized.status_code, 400)
        self.assertIn("cover limit", oversized.json()["detail"])

    # -- submission -------------------------------------------------------

    def test_submission_gates_and_read_only_submitted_state(self):
        owner = ApiClient()
        bare = self.create_draft(owner, title="", body="")
        blocked = owner.request("POST", f"/internal-content/{bare['id']}/submit")
        self.assertEqual(blocked.status_code, 400)
        detail = blocked.json()["detail"]
        for expected in ("title", "story body", "cover"):
            self.assertIn(expected, detail)

        without_cover = self.create_draft(owner)
        missing_cover = owner.request("POST", f"/internal-content/{without_cover['id']}/submit")
        self.assertEqual(missing_cover.status_code, 400)
        self.assertIn("cover", missing_cover.json()["detail"])

        submitted = self.create_draft(owner)
        self.upload_cover(owner, submitted["id"], build_image(1920, 1080))
        done = owner.request("POST", f"/internal-content/{submitted['id']}/submit")
        self.assertEqual(done.status_code, 200)
        final = done.json()
        self.assertEqual(final["status"], "submitted")
        self.assertTrue(final["submitted_at"])

        again = owner.request("POST", f"/internal-content/{submitted['id']}/submit")
        self.assertEqual(again.status_code, 200)

        edit_attempt = owner.request("PUT", f"/internal-content/{submitted['id']}", json_body={
            "title": "Should not apply",
            "body": "Should not apply.",
        })
        self.assertEqual(edit_attempt.status_code, 400)
        delete_attempt = owner.request("DELETE", f"/internal-content/{submitted['id']}")
        self.assertEqual(delete_attempt.status_code, 200)
        self.assertEqual(
            owner.request("GET", f"/internal-content/{submitted['id']}").status_code,
            404,
        )

    # -- editorial review (privileged) -------------------------------------

    EDITOR_KEY = "editor-suite-key-77"

    def _submit_story(self, client, **fields):
        record = self.create_draft(client, **fields)
        self.upload_cover(client, record["id"], build_image(1920, 1080))
        done = client.request("POST", f"/internal-content/{record['id']}/submit")
        self.assertEqual(done.status_code, 200)
        return record

    def test_editor_endpoints_require_the_internal_editor_key(self):
        owner = ApiClient()
        record = self._submit_story(owner)

        missing = owner.request("GET", "/internal-content/review")
        self.assertEqual(missing.status_code, 403)
        wrong = owner.request("GET", "/internal-content/review", headers={"x-editor-key": "wrong-key"})
        self.assertEqual(wrong.status_code, 403)

        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            listing = owner.request(
                "GET", "/internal-content/review", headers={"x-editor-key": self.EDITOR_KEY}
            )
        self.assertEqual(listing.status_code, 200)
        ids = [item["id"] for item in listing.json()["items"]]
        self.assertIn(record["id"], ids)

    def test_publish_is_instant_public_and_notifies_only_the_author(self):
        owner, outsider = ApiClient(), ApiClient()
        record = self._submit_story(owner, title="Facility update")

        hidden = outsider.request("GET", f"/internal-content/{record['id']}/cover")
        self.assertEqual(hidden.status_code, 404)

        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            published = outsider.request(
                "POST", f"/internal-content/{record['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
        self.assertEqual(published.status_code, 200)
        body = published.json()
        self.assertEqual(body["status"], "published")
        self.assertTrue(body["published_at"])

        feed = outsider.request("GET", "/internal-content/published").json()["items"]
        self.assertTrue(any(item["id"] == record["id"] for item in feed))

        public_reader = outsider.request(
            "GET", f"/internal-content/published/{record['id']}"
        )
        self.assertEqual(public_reader.status_code, 200)
        self.assertEqual(public_reader.json()["id"], record["id"])

        visible = outsider.request("GET", f"/internal-content/{record['id']}/cover")
        self.assertEqual(visible.status_code, 200)

        inbox = owner.request("GET", "/internal-content/notifications").json()
        self.assertEqual(inbox["unread"], 1)
        self.assertEqual(inbox["items"][0]["kind"], "published")
        self.assertEqual(inbox["items"][0]["record_id"], record["id"])

        other = ApiClient().request("GET", "/internal-content/notifications").json()
        self.assertEqual(other["unread"], 0)

        marked = owner.request(
            "POST", "/internal-content/notifications/read",
            json_body={"ids": [inbox["items"][0]["id"]]},
        )
        self.assertEqual(marked.json()["unread"], 0)

    def test_author_can_permanently_delete_published_contribution(self):
        owner = ApiClient()
        record = self._submit_story(owner, title="Remove published story")
        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            published = owner.request(
                "POST", f"/internal-content/{record['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
        self.assertEqual(published.status_code, 200)
        self.assertEqual(
            owner.request("GET", f"/internal-content/published/{record['id']}").status_code,
            200,
        )
        self.assertEqual(
            owner.request("GET", "/internal-content/notifications").json()["unread"],
            1,
        )

        deleted = owner.request("DELETE", f"/internal-content/{record['id']}")
        self.assertEqual(deleted.status_code, 200)
        self.assertEqual(deleted.json(), {"deleted": True})
        self.assertEqual(
            owner.request("GET", f"/internal-content/published/{record['id']}").status_code,
            404,
        )
        self.assertEqual(owner.request("GET", "/internal-content/mine").json()["items"], [])
        self.assertEqual(
            owner.request("GET", "/internal-content/notifications").json()["items"],
            [],
        )

    def test_author_can_withdraw_submission_and_resume_editing(self):
        owner = ApiClient()
        record = self._submit_story(owner, title="Withdrawable story")

        withdrawn = owner.request("POST", f"/internal-content/{record['id']}/withdraw")
        self.assertEqual(withdrawn.status_code, 200)
        self.assertEqual(withdrawn.json()["status"], "withdrawn")
        self.assertTrue(withdrawn.json()["withdrawn_at"])

        revised = owner.request("PUT", f"/internal-content/{record['id']}", json_body={
            "title": "Revised after withdrawal",
            "body": "The author can safely continue editing this withdrawn submission.",
        })
        self.assertEqual(revised.status_code, 200)
        self.assertEqual(revised.json()["title"], "Revised after withdrawal")

    def test_editor_archive_and_restore_preserve_published_record(self):
        owner = ApiClient()
        record = self._submit_story(owner, title="Archive lifecycle story")
        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            published = owner.request(
                "POST", f"/internal-content/{record['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
            self.assertEqual(published.status_code, 200)
            archived = owner.request(
                "POST", f"/internal-content/{record['id']}/archive",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
            self.assertEqual(archived.status_code, 200)
            self.assertEqual(archived.json()["status"], "archived")
            self.assertEqual(owner.request("GET", f"/internal-content/published/{record['id']}").status_code, 404)

            restored = owner.request(
                "POST", f"/internal-content/{record['id']}/restore",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
        self.assertEqual(restored.status_code, 200)
        self.assertEqual(restored.json()["status"], "published")
        self.assertEqual(owner.request("GET", f"/internal-content/published/{record['id']}").status_code, 200)

    def test_expired_publication_disappears_without_deleting_author_record(self):
        owner = ApiClient()
        record = self._submit_story(
            owner,
            title="Expired announcement",
            content_type="announcement",
            expires_at="2020-01-01T00:00:00Z",
        )
        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            published = owner.request(
                "POST", f"/internal-content/{record['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
        self.assertEqual(published.status_code, 200)
        self.assertEqual(published.json()["status"], "published")
        self.assertEqual(owner.request("GET", f"/internal-content/published/{record['id']}").status_code, 404)
        mine = {item["id"]: item for item in owner.request("GET", "/internal-content/mine").json()["items"]}
        self.assertEqual(mine[record["id"]]["status"], "published")

    def test_public_reader_rejects_drafts_and_missing_records(self):
        owner = ApiClient()
        draft = self.create_draft(owner, title="Private draft", body="This remains private.")
        self.assertEqual(
            owner.request("GET", f"/internal-content/published/{draft['id']}").status_code,
            404,
        )
        self.assertEqual(
            owner.request("GET", "/internal-content/published/does-not-exist").status_code,
            404,
        )

    def test_request_changes_lets_the_author_revise_and_resubmit(self):
        owner = ApiClient()
        record = self._submit_story(owner)

        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            empty = owner.request(
                "POST", f"/internal-content/{record['id']}/changes",
                headers={"x-editor-key": self.EDITOR_KEY}, json_body={"note": ""},
            )
            self.assertEqual(empty.status_code, 400)

            decided = owner.request(
                "POST", f"/internal-content/{record['id']}/changes",
                headers={"x-editor-key": self.EDITOR_KEY},
                json_body={"note": "Tighten the summary."},
            )
        self.assertEqual(decided.status_code, 200)
        body = decided.json()
        self.assertEqual(body["status"], "needs_changes")
        self.assertEqual(body["review_note"], "Tighten the summary.")

        revised = owner.request("PUT", f"/internal-content/{record['id']}", json_body={
            "title": "Covered story",
            "summary": "A tighter summary.",
            "body": "A body long enough to submit safely.",
        })
        self.assertEqual(revised.status_code, 200)
        resubmitted = owner.request("POST", f"/internal-content/{record['id']}/submit")
        self.assertEqual(resubmitted.status_code, 200)
        self.assertEqual(resubmitted.json()["status"], "submitted")

        inbox = owner.request("GET", "/internal-content/notifications").json()
        self.assertTrue(any(item["kind"] == "changes" for item in inbox["items"]))

    def test_reject_archives_and_blocks_further_decisions(self):
        owner = ApiClient()
        record = self._submit_story(owner)

        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            rejected = owner.request(
                "POST", f"/internal-content/{record['id']}/reject",
                headers={"x-editor-key": self.EDITOR_KEY}, json_body={"note": "Duplicate story."},
            )
        self.assertEqual(rejected.status_code, 200)
        self.assertEqual(rejected.json()["status"], "archived")

        feed = owner.request("GET", "/internal-content/published").json()["items"]
        self.assertFalse(any(item["id"] == record["id"] for item in feed))

        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            again = owner.request(
                "POST", f"/internal-content/{record['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
        self.assertEqual(again.status_code, 400)

        mine = [item["id"] for item in owner.request("GET", "/internal-content/mine").json()["items"]]
        self.assertIn(record["id"], mine)
        inbox = owner.request("GET", "/internal-content/notifications").json()
        self.assertTrue(any(item["kind"] == "rejected" for item in inbox["items"]))

    def test_publishing_requires_submitted_status(self):
        owner = ApiClient()
        draft = self.create_draft(owner)
        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            early = owner.request(
                "POST", f"/internal-content/{draft['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
        self.assertEqual(early.status_code, 400)

    def test_review_unlock_sets_httponly_cookie_session_and_lock_revokes_it(self):
        owner = ApiClient()
        record = self._submit_story(owner)

        wrong = owner.request("POST", "/internal-content/review/unlock", json_body={"key": "nope"})
        self.assertEqual(wrong.status_code, 403)

        # The whole session lives under one stable server configuration.
        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            unlocked = owner.request(
                "POST", "/internal-content/review/unlock", json_body={"key": self.EDITOR_KEY}
            )
            self.assertEqual(unlocked.status_code, 200)

            # The session cookie rides along; no key header needed anymore.
            listing = owner.request("GET", "/internal-content/review")
            self.assertEqual(listing.status_code, 200)
            ids = [item["id"] for item in listing.json()["items"]]
            self.assertIn(record["id"], ids)
            set_cookie = unlocked.headers.get("set-cookie", "")
            self.assertIn("HttpOnly", set_cookie)

            locked = owner.request("POST", "/internal-content/review/lock")
        self.assertEqual(locked.status_code, 200)
        denied = owner.request("GET", "/internal-content/review")
        self.assertEqual(denied.status_code, 403)

    # -- leadership channel -------------------------------------------------

    def test_leadership_publish_retires_the_previous_vision(self):
        owner = ApiClient()

        def submit_vision(title):
            record = self.create_draft(owner, title=title, content_type="leadership", category="Leadership")
            self.upload_cover(owner, record["id"], build_image(1920, 1080))
            done = owner.request("POST", f"/internal-content/{record['id']}/submit")
            self.assertEqual(done.status_code, 200)
            return record

        first = submit_vision("Vision of the quarter")
        second = submit_vision("Vision of the new quarter")

        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            published_first = owner.request(
                "POST", f"/internal-content/{first['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
            self.assertEqual(published_first.status_code, 200)
            feed_one = owner.request("GET", "/internal-content/published").json()["items"]
            leadership_one = [item for item in feed_one if item["content_type"] == "leadership"]
            self.assertEqual([item["id"] for item in leadership_one], [first["id"]])

            published_second = owner.request(
                "POST", f"/internal-content/{second['id']}/publish",
                headers={"x-editor-key": self.EDITOR_KEY},
            )
            self.assertEqual(published_second.status_code, 200)

        feed_two = owner.request("GET", "/internal-content/published").json()["items"]
        leadership_two = [item for item in feed_two if item["content_type"] == "leadership"]
        self.assertEqual([item["id"] for item in leadership_two], [second["id"]])

        mine = {item["id"]: item for item in owner.request("GET", "/internal-content/mine").json()["items"]}
        self.assertEqual(mine[first["id"]]["status"], "archived")
        self.assertEqual(mine[second["id"]]["status"], "published")

    def test_unknown_content_type_is_rejected(self):
        owner = ApiClient()
        response = owner.request("POST", "/internal-content/drafts", json_body={
            "title": "Strange type",
            "content_type": "newsletter",
        })
        self.assertEqual(response.status_code, 400)

    def test_announcements_publish_and_accumulate_without_retirement(self):
        owner = ApiClient()

        def submit_announcement(title):
            record = self.create_draft(owner, title=title, content_type="announcement", category="Announcement")
            self.upload_cover(owner, record["id"], build_image(1920, 1080))
            done = owner.request("POST", f"/internal-content/{record['id']}/submit")
            self.assertEqual(done.status_code, 200)
            return record

        first = submit_announcement("Town hall on Friday")
        second = submit_announcement("New campus access badges")

        with patch.dict(os.environ, {"INTERNAL_EDITOR_KEY": self.EDITOR_KEY}):
            for record in (first, second):
                published = owner.request(
                    "POST", f"/internal-content/{record['id']}/publish",
                    headers={"x-editor-key": self.EDITOR_KEY},
                )
                self.assertEqual(published.status_code, 200)

        feed = owner.request("GET", "/internal-content/published").json()["items"]
        announcements = [item for item in feed if item["content_type"] == "announcement"]
        # Unlike leadership, announcements are a stream: every published one stays live.
        self.assertEqual(
            sorted(item["id"] for item in announcements),
            sorted([first["id"], second["id"]]),
        )

    def test_announcements_submit_without_a_cover_and_imports_can_route_to_them(self):
        owner = ApiClient()

        # Text-first notice: no cover attached at all.
        bare = self.create_draft(owner, title="Cafeteria hours change", body="The cafe now opens at eight.", content_type="announcement", category="Announcement")
        submitted = owner.request("POST", f"/internal-content/{bare['id']}/submit")
        self.assertEqual(submitted.status_code, 200)
        self.assertEqual(submitted.json()["status"], "submitted")

        # A PDF import can be routed straight into the announcement channel.
        imported = owner.upload("POST", "/internal-content/import", {
            "document": ("notice.pdf", build_pdf(["Quarterly all-hands notice for every regional office this Friday"]), "application/pdf"),
            "content_type": (None, "announcement", None),
        })
        self.assertEqual(imported.status_code, 200)
        record = imported.json()
        self.assertEqual(record["content_type"], "announcement")
        self.assertIn("Quarterly all-hands notice", record["body"])

        submitted_notice = owner.request("POST", f"/internal-content/{record['id']}/submit")
        self.assertEqual(submitted_notice.status_code, 200)

        # Stories keep their cover requirement.
        story_owner = ApiClient()
        storyless = self.create_draft(story_owner, title="A story without a cover", body="Long enough body text goes here.")
        blocked = story_owner.request("POST", f"/internal-content/{storyless['id']}/submit")
        self.assertEqual(blocked.status_code, 400)
        self.assertIn("cover", blocked.json()["detail"])

    def test_published_endpoint_only_returns_published_status(self):
        owner = ApiClient()
        record = self.create_draft(owner)
        self.upload_cover(owner, record["id"], build_image(1600, 900))
        owner.request("POST", f"/internal-content/{record['id']}/submit")

        published = owner.request("GET", "/internal-content/published").json()["items"]
        self.assertEqual(published, [])

        future_record = self.create_draft(ApiClient(), title="Already published", body="Legacy published content.")
        with internal_storage.mutation_lock:
            items = internal_storage.load_records()
            items[future_record["id"]]["status"] = "published"
            items[future_record["id"]]["published_at"] = items[future_record["id"]]["updated_at"]
            internal_storage.write_records(items)

        published_now = owner.request("GET", "/internal-content/published").json()["items"]
        self.assertEqual([item["id"] for item in published_now], [future_record["id"]])
        self.assertFalse(any(item["id"] == record["id"] for item in published_now))

    def test_delete_draft_removes_only_its_own_files(self):
        owner = ApiClient()
        keep = self.create_draft(owner)
        self.upload_cover(owner, keep["id"], build_image(1600, 900), name="keep.png")
        remove = self.create_draft(owner)
        self.upload_cover(owner, remove["id"], build_image(1600, 900), name="remove.png")
        imported = owner.upload("POST", "/internal-content/import", {
            "document": ("brief.docx", build_docx(["Briefing body text with plenty of readable characters for the deletion flow."]), ""),
        }).json()

        covers_before = {path.name for path in Path(internal_storage.COVERS_DIR).glob("*")}
        originals_before = {path.name for path in Path(internal_storage.ORIGINALS_DIR).glob("*")}
        self.assertEqual(len(covers_before), 2)
        self.assertEqual(len(originals_before), 1)

        removed_cover_file = internal_storage.load_records()[remove["id"]]["cover"]["file"]
        removed_original_file = imported["source_document"]["file"]
        kept_cover_file = internal_storage.load_records()[keep["id"]]["cover"]["file"]

        deleted = owner.request("DELETE", f"/internal-content/{remove['id']}")
        self.assertEqual(deleted.status_code, 200)
        deleted_import = owner.request("DELETE", f"/internal-content/{imported['id']}")
        self.assertEqual(deleted_import.status_code, 200)

        covers_after = {path.name for path in Path(internal_storage.COVERS_DIR).glob("*")}
        originals_after = {path.name for path in Path(internal_storage.ORIGINALS_DIR).glob("*")}
        self.assertEqual(covers_after, {kept_cover_file})
        self.assertNotIn(removed_cover_file, covers_after)
        self.assertNotIn(removed_original_file, originals_after)


if __name__ == "__main__":
    unittest.main()
